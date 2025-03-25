from flask import Flask, render_template, request, jsonify, send_file
from quiz_generator import generate_quiz
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import tempfile
import random
import re

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-key-here')

# Load IRPCS rules from JSON file
with open('irpcs_rules.json', 'r', encoding='utf-8') as f:
    IRPCS_RULES = json.load(f)

def create_word_document(questions, title="IRPCS Quiz", show_answers=False):
    try:
        print(f"Creating Word document with {len(questions)} questions")
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        doc.add_paragraph()  # Add spacing
        
        # Validate questions
        valid_questions = []
        for q in questions:
            if not isinstance(q, dict) or 'question' not in q or 'type' not in q:
                print(f"Skipping invalid question: {q}")
                continue
                
            if q['type'] == 'Multiple Choice' and ('choices' not in q or not q['choices']):
                print(f"Skipping invalid multiple choice question without choices: {q}")
                continue
                
            valid_questions.append(q)
            
        if not valid_questions:
            # Add error message to document if no valid questions
            error_paragraph = doc.add_paragraph("No valid questions could be generated. Please try again.")
            error_paragraph.style = 'Intense Quote'
            return doc
        
        # Add questions
        for i, question in enumerate(valid_questions, 1):
            try:
                # Question text and number
                question_paragraph = doc.add_paragraph()
                question_paragraph.add_run(f"{i}. ").bold = True
                
                # Handle marks and format the question appropriately
                if 'marks' in question:
                    # Check if marks are already included in the question text
                    if str(question['marks']) + " Marks" not in question['question'] and str(question['marks']) + " Mark" not in question['question']:
                        # If the question contains line breaks for parts, place marks at the end
                        if '\n' in question['question'] or '<br>' in question['question']:
                            question_text = question['question']
                        else:
                            # Otherwise add marks after the question
                            mark_text = f" {question['marks']} Mark" if question['marks'] == 1 else f" {question['marks']} Marks"
                            if question['question'].endswith(mark_text):
                                question_text = question['question']
                            else:
                                question_text = question['question']
                    else:
                        question_text = question['question']
                else:
                    question_text = question['question']
                
                # Replace any HTML line breaks with actual line breaks
                question_text = question_text.replace('<br>', '\n')
                
                # Add each line of the question, handling multi-part questions
                for line in question_text.split('\n'):
                    # Bold the part identifiers (a), b), etc.)
                    if line.strip().startswith('a)') or line.strip().startswith('b)') or line.strip().startswith('c)') or line.strip().startswith('d)') or line.strip().startswith('e)') or line.strip().startswith('f)'):
                        part_run = question_paragraph.add_run(line[:2])
                        part_run.bold = True
                        question_paragraph.add_run(line[2:])
                    else:
                        question_paragraph.add_run(line)
                    question_paragraph.add_run('\n')
                
                # Add choices for multiple choice questions
                if question['type'] == 'Multiple Choice' and 'choices' in question:
                    for idx, choice in enumerate(question['choices']):
                        choice_paragraph = doc.add_paragraph()
                        choice_paragraph.style = 'List Bullet'
                        letter = chr(97 + idx)  # a, b, c, d...
                        choice_paragraph.add_run(f"{letter}) {choice}")
                
                # Add true/false options
                elif question['type'] == 'True/False':
                    true_false_paragraph = doc.add_paragraph()
                    true_false_paragraph.add_run("☐ True      ☐ False")
                
                # Add answer space for fill in the gap and written answer
                elif question['type'] in ['Fill in the Gap', 'Written Answer']:
                    # For Fill in the Gap, add a single answer line
                    if question['type'] == 'Fill in the Gap':
                        doc.add_paragraph("Answer: _________________________________________________")
                    # For Written Answer, add answer lines based on parts
                    else:
                        # Check if the question has lettered parts
                        if 'a)' in question['question']:
                            parts = ['a)', 'b)', 'c)', 'd)', 'e)', 'f)']
                            existing_parts = [p for p in parts if p in question['question']]
                            
                            for part in existing_parts:
                                answer_paragraph = doc.add_paragraph()
                                answer_paragraph.add_run(f"{part} ").bold = True
                                answer_paragraph.add_run("_________________________________________________")
                        else:
                            # If no lettered parts, add a generic answer line
                            doc.add_paragraph("Answer: _________________________________________________")
                
                # Add rule reference if available
                if question.get('rule_reference'):
                    rule_paragraph = doc.add_paragraph()
                    rule_paragraph.style = 'Intense Quote'
                    rule_paragraph.add_run(f"Rule Reference: {question['rule_reference']}")
                
                # Add mark allocation (if not already included in the question)
                if 'marks' in question and "Marks" not in question['question'] and "Mark" not in question['question']:
                    marks_paragraph = doc.add_paragraph()
                    marks_text = f"({question['marks']} mark)" if question['marks'] == 1 else f"({question['marks']} marks)"
                    marks_paragraph.add_run(marks_text).italic = True
                    marks_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Add answer if showing answers
                if show_answers and 'answer' in question:
                    answer_paragraph = doc.add_paragraph()
                    answer_paragraph.style = 'Intense Quote'
                    answer_paragraph.add_run("Answer: ").bold = True
                    
                    # Handle multi-part answers with line breaks
                    if '\n' in question['answer']:
                        # First add the "Answer:" prefix
                        answer_lines = question['answer'].split('\n')
                        for line in answer_lines:
                            if line.strip():
                                # Check if this line starts with a part identifier (a), b), etc.)
                                if line.strip().startswith('a)') or line.strip().startswith('b)') or line.strip().startswith('c)') or line.strip().startswith('d)') or line.strip().startswith('e)') or line.strip().startswith('f)'):
                                    part_run = answer_paragraph.add_run(f"\n{line[:2]}")
                                    part_run.bold = True
                                    answer_paragraph.add_run(line[2:])
                                else:
                                    answer_paragraph.add_run(f"\n{line}")
                    else:
                        answer_paragraph.add_run(question['answer'])
                    
                    # Add explanation for True/False questions if available
                    if question['type'] == 'True/False' and 'explanation' in question:
                        explanation_paragraph = doc.add_paragraph()
                        explanation_paragraph.style = 'Quote'
                        explanation_paragraph.add_run("Explanation: ").italic = True
                        explanation_paragraph.add_run(question['explanation']).italic = True
                
                doc.add_paragraph()  # Add spacing between questions
            except Exception as q_error:
                print(f"Error adding question {i}: {str(q_error)}")
                # Add error note instead of skipping completely
                error_paragraph = doc.add_paragraph()
                error_paragraph.add_run(f"{i}. ").bold = True
                error_paragraph.add_run("Error processing this question. Please try again.")
                doc.add_paragraph()
        
        return doc
    except Exception as e:
        print(f"Error in create_word_document: {str(e)}")
        # Create a minimal document with error message
        doc = Document()
        doc.add_paragraph("Error generating quiz. Please try again.").style = 'Intense Quote'
        return doc

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/rules')
def rules():
    return render_template('rules.html', IRPCS_RULES=IRPCS_RULES)

@app.route('/policy')
def policy():
    return render_template('policy.html')

@app.route('/generate', methods=['GET', 'POST'])
def generate():
    if request.method == 'GET':
        return render_template('generate.html')
    elif request.method == 'POST':
        try:
            print("Received generate request")
            
            # Check content type header
            content_type = request.headers.get('Content-Type', '')
            if 'application/json' not in content_type:
                print(f"Invalid content type: {content_type}")
                return jsonify({
                    'success': False,
                    'message': f"415 Unsupported Media Type: Expected 'application/json', got '{content_type}'"
                }), 415
                
            # Parse JSON data
            data = request.get_json(silent=True)
            if not data:
                print("Failed to parse JSON data")
                return jsonify({
                    'success': False,
                    'message': "400 Bad Request: Invalid JSON data"
                }), 400
            
            print(f"Processing quiz request: {data}")
            
            preview_only = data.get('preview_only', False)
            show_answers = data.get('show_answers', False)
            
            # Check if we're using pre-generated questions
            if 'questions' in data and isinstance(data['questions'], list) and data['questions']:
                print(f"Using {len(data['questions'])} pre-generated questions")
                questions = data['questions']
            else:
                # Generate new questions
                num_questions = min(int(data.get('num_questions', 10)), 50)
                difficulty = data.get('difficulty', 'mixed')
                question_types = data.get('question_types', ['Multiple Choice', 'Fill in the Gap', 'Written Answer', 'True/False'])
                
                print(f"Generating {num_questions} questions, difficulty={difficulty}, types={question_types}, show_answers={show_answers}")
                
                questions = []
                for _ in range(num_questions):
                    question = generate_question(difficulty, question_types)
                    questions.append(question)
                
                print(f"Generated {len(questions)} questions")
            
            # If preview only, return JSON data
            if preview_only:
                return jsonify({
                    'success': True,
                    'questions': questions
                })
            
            # Otherwise, create and send document for download
            doc = create_word_document(questions, title="IRPCS Quiz", show_answers=show_answers)
            
            # Save to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp_path = tmp.name
                print(f"Saving document to {tmp_path}")
                doc.save(tmp_path)
                
                # Return the file
                try:
                    print(f"Sending file {tmp_path}")
                    return send_file(
                        tmp_path,
                        as_attachment=True,
                        download_name=f'irpcs_quiz_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                except Exception as file_error:
                    print(f"Error sending file: {str(file_error)}")
                    # Clean up temp file
                    try:
                        os.unlink(tmp_path)
                    except:
                        pass
                    raise

        except Exception as e:
            import traceback
            error_traceback = traceback.format_exc()
            print(f"Error in generate: {str(e)}\n{error_traceback}")
            return jsonify({
                'success': False,
                'message': f"Error generating quiz: {str(e)}"
            }), 500

@app.route('/api/edit', methods=['POST'])
def edit_question():
    try:
        data = request.get_json()
        index = data.get('index')
        question = data.get('question')

        # Validate question data
        if not question.get('question') or not question.get('answer'):
            raise ValueError("Question and answer are required")

        return jsonify({
            'success': True,
            'message': 'Question updated successfully'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        })

@app.route('/download')
def download_quiz():
    try:
        show_answers = request.args.get('show_answers', 'false').lower() == 'true'
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run('IRPCS Quiz')
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()
            
            # Generate questions
            questions = []
            for _ in range(10):  # Generate 10 questions
                question = generate_question('mixed', ['Multiple Choice', 'Fill in the Gap', 'Written Answer'])
                questions.append(question)
            
            # Add questions
            for i, question in enumerate(questions, 1):
                # Question number and text
                question_paragraph = doc.add_paragraph()
                question_paragraph.add_run(f"{i}. ").bold = True
                question_paragraph.add_run(question['question'])
                
                # Add choices for multiple choice questions
                if question['type'] == 'Multiple Choice':
                    for choice in question['choices']:
                        doc.add_paragraph(choice, style='List Bullet')
                
                # Add answer space for fill in the gap and written answer
                elif question['type'] in ['Fill in the Gap', 'Written Answer']:
                    doc.add_paragraph("Answer: _________________________________________________")
                
                # Add rule reference
                doc.add_paragraph(f"Rule Reference: {question['rule_reference']}", style='Intense Quote')
                
                # Add answer if showing answers
                if show_answers:
                    answer_paragraph = doc.add_paragraph()
                    answer_paragraph.add_run("Answer: ").bold = True
                    if question['type'] == 'Multiple Choice':
                        answer_paragraph.add_run(question['choices'][question['correct_index']])
                    else:
                        answer_paragraph.add_run(question['answer'])
                    answer_paragraph.style = 'Intense Quote'
                
                doc.add_paragraph()  # Add spacing between questions
            
            doc.save(tmp.name)
            
            return send_file(
                tmp.name,
                as_attachment=True,
                download_name=f'irpcs_quiz_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            )
            
    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        })

@app.route('/generate/question', methods=['POST'])
def regenerate_question():
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'message': "400 Bad Request: Invalid JSON data"
            }), 400
            
        difficulty = data.get('difficulty', 'mixed')
        question_types = [data.get('question_type', 'Multiple Choice')]
        
        # Generate a new question
        question = generate_question(difficulty, question_types)
        
        return jsonify({
            'success': True,
            'question': question
        })
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"Error in regenerate_question: {str(e)}\n{error_traceback}")
        return jsonify({
            'success': False,
            'message': f"Error generating question: {str(e)}"
        }), 500

def generate_question(difficulty, question_types):
    # Get a random question type if multiple are available
    question_type = random.choice(question_types)
    
    # Check that IRPCS_RULES contains 'rules' key and it's not empty
    if 'rules' not in IRPCS_RULES or not IRPCS_RULES['rules']:
        print("ERROR: Invalid IRPCS_RULES structure - missing or empty 'rules' key")
        raise ValueError("Invalid rules data structure")
    
    # Filter rules to only select from rules 1-35
    valid_rules = [rule for rule in IRPCS_RULES['rules'] if rule.get('id') and int(rule['id']) >= 1 and int(rule['id']) <= 35]
    
    if not valid_rules:
        print("ERROR: No rules found in range 1-35")
        raise ValueError("No valid rules found in specified range")
    
    # Get a random rule from the filtered list
    rule = random.choice(valid_rules)
    
    # Validate rule has required fields
    if not rule or 'id' not in rule or 'title' not in rule or 'content' not in rule:
        print(f"ERROR: Invalid rule structure: {rule}")
        raise ValueError(f"Invalid rule data structure: {rule}")
    
    # Generate exam-style questions based on specific rules
    if question_type == 'Multiple Choice':
        return generate_exam_multiple_choice(rule, difficulty)
    elif question_type == 'Fill in the Gap':
        return generate_exam_fill_blank(rule, difficulty)
    elif question_type == 'True/False':
        return generate_true_false(rule, difficulty)
    else:  # 'Written Answer'
        return generate_exam_written_answer(rule, difficulty)

def generate_exam_multiple_choice(rule, difficulty):
    # Create exam-style multiple choice questions specifically referencing rules
    rule_questions = {
        # Rule-specific questions - each rule ID should have one or more question templates
        '1': [
            {
                'question': f"From Rule {rule['id']}, where do the Regulations apply?",
                'options': [
                    "To all vessels upon the high seas and in all waters connected therewith navigable by seagoing vessels",
                    "Only to vessels in international waters",
                    "Only to commercial vessels",
                    "Only to vessels over 20 meters in length"
                ],
                'correct_index': 0,
                'marks': 2
            }
        ],
        '3': [
            {
                'question': f"What is the definition of the word 'vessel'?",
                'options': [
                    "Any motorized watercraft used for transportation",
                    "Every description of watercraft, including non-displacement craft, WIG craft and seaplanes, used or capable of being used as a means of transportation on water",
                    "A ship of substantial size used for commercial purposes",
                    "Any watercraft powered by mechanical means"
                ],
                'correct_index': 1,
                'marks': 3
            },
            {
                'question': f"What is the definition of 'vessel engaged in fishing'?",
                'options': [
                    "Any vessel fishing with nets, lines, trawls or other fishing apparatus which restrict maneuverability",
                    "Any vessel collecting fish from the water",
                    "Any vessel using equipment to catch fish",
                    "Any vessel with fishing equipment visible on deck"
                ],
                'correct_index': 0,
                'marks': 2
            },
            {
                'question': f"What is the definition of a 'vessel not under command'?",
                'options': [
                    "A vessel without a crew",
                    "A vessel which through some exceptional circumstance is unable to maneuver as required by these Rules",
                    "A vessel without a captain",
                    "A vessel that is anchored or moored"
                ],
                'correct_index': 1,
                'marks': 2
            },
            {
                'question': f"What is the definition of the word 'underway'?",
                'options': [
                    "A vessel moving through the water",
                    "A vessel making way through the water using its engine",
                    "A vessel that is not at anchor, or made fast to the shore, or aground",
                    "A vessel traveling at a speed above 5 knots"
                ],
                'correct_index': 2,
                'marks': 2
            },
            {
                'question': f"The term 'power-driven vessel' means:",
                'options': [
                    "Any vessel propelled by machinery",
                    "Any vessel with an engine, even if not using it",
                    "Only vessels with twin engines",
                    "Any vessel larger than 12 meters"
                ],
                'correct_index': 0,
                'marks': 2
            }
        ],
        '5': [
            {
                'question': f"What is required to maintain a proper look-out?",
                'options': [
                    "Using radar equipment only",
                    "Every vessel shall at all times maintain a proper look-out by sight and hearing",
                    "Monitoring radio communications",
                    "Scanning the horizon with binoculars every fifteen minutes"
                ],
                'correct_index': 1,
                'marks': 2
            }
        ],
        '6': [
            {
                'question': f"Which of the following is a factor that should be taken into account when determining a safe speed?",
                'options': [
                    "The traffic density",
                    "The age of the vessel",
                    "The nationality of the vessel",
                    "The color of the vessel"
                ],
                'correct_index': 0,
                'marks': 2
            }
        ],
        '7': [
            {
                'question': f"When are vessels deemed to be in sight of one another?",
                'options': [
                    "When one can be observed visually from the other",
                    "When they are within 10 nautical miles of each other",
                    "When they can communicate by VHF radio",
                    "When they can be detected on radar"
                ],
                'correct_index': 0,
                'marks': 2
            },
            {
                'question': f"What action should be taken if there is any doubt about risk of collision?",
                'options': [
                    "Maintain course and speed",
                    "Reduce speed by half",
                    "Assume that risk of collision exists",
                    "Call the other vessel on VHF radio"
                ],
                'correct_index': 2,
                'marks': 2
            }
        ],
        '8': [
            {
                'question': f"What do the rules state about any action taken to avoid collision?",
                'options': [
                    "It shall be small and gradual",
                    "It shall be positive, made in ample time and with due regard to good seamanship",
                    "It shall be made only after communication with the other vessel",
                    "It shall be made only by altering course, never by changing speed"
                ],
                'correct_index': 1,
                'marks': 2
            }
        ],
        '9': [
            {
                'question': f"In what part of a narrow channel should you navigate?",
                'options': [
                    "In the center of the channel",
                    "As near to the outer limit which lies on her starboard side as is safe and practicable",
                    "As close as possible to the port side",
                    "In the deepest part of the channel regardless of position"
                ],
                'correct_index': 1,
                'marks': 2
            },
            {
                'question': f"Are you within your rights if you anchor in a narrow channel in a case of emergency?",
                'options': [
                    "Yes",
                    "No",
                    "Only with permission from authorities",
                    "Only during daylight hours"
                ],
                'correct_index': 0,
                'marks': 1
            }
        ],
        '12': [
            {
                'question': f"If two sailing vessels were approaching one another on a steady bearing, which would give way when each has the wind on a different side?",
                'options': [
                    "The vessel which has the wind on the port side",
                    "The vessel which has the wind on the starboard side",
                    "The vessel with more sail area",
                    "The vessel that is moving faster"
                ],
                'correct_index': 0,
                'marks': 2
            }
        ],
        '13': [
            {
                'question': f"When is a vessel deemed to be an overtaking vessel?",
                'options': [
                    "When it is approaching another vessel from a direction more than 22.5° abaft the beam",
                    "When it is approaching another vessel from ahead",
                    "When it is passing within 1 nautical mile of another vessel",
                    "When it is moving at a greater speed than the vessel ahead"
                ],
                'correct_index': 0,
                'marks': 2
            },
            {
                'question': f"A vessel being overtaken by another vessel should:",
                'options': [
                    "Increase speed to maintain separation",
                    "Alter course to port to assist the overtaking vessel",
                    "Keep her course and speed",
                    "Sound one prolonged blast"
                ],
                'correct_index': 2,
                'marks': 2
            }
        ],
        '14': [
            {
                'question': f"Which vessel has priority in a head-on situation?",
                'options': [
                    "The larger vessel",
                    "The faster vessel",
                    "Neither - both should alter course to starboard",
                    "The vessel with radar"
                ],
                'correct_index': 2,
                'marks': 2
            },
            {
                'question': f"If you see ahead of you a power-driven vessel on, or nearly on a reciprocal course so as to involve risk of collision what action would you take?",
                'options': [
                    "Alter course to starboard",
                    "Alter course to port",
                    "Maintain course and speed",
                    "Reduce speed but maintain course"
                ],
                'correct_index': 0,
                'marks': 2
            }
        ],
        '15': [
            {
                'question': f"In a crossing situation between two power-driven vessels, which vessel should keep out of the way?",
                'options': [
                    "The vessel which has the other on her starboard side",
                    "The vessel which has the other on her port side",
                    "The vessel that is traveling faster",
                    "The smaller vessel"
                ],
                'correct_index': 0,
                'marks': 2
            },
            {
                'question': f"Which vessel has the right of way in a crossing situation?",
                'options': [
                    "The larger vessel",
                    "The vessel on the starboard side",
                    "The vessel on the port side",
                    "The faster vessel"
                ],
                'correct_index': 1,
                'marks': 2
            }
        ],
        '18': [
            {
                'question': f"A sailing vessel approaching a power-driven vessel should:",
                'options': [
                    "Always alter course to starboard",
                    "Keep her course",
                    "Give way",
                    "Show day shapes"
                ],
                'correct_index': 1,
                'marks': 2
            },
            {
                'question': f"Which of the following vessels must a sailing vessel keep out of the way of when underway?",
                'options': [
                    "A power-driven vessel",
                    "A vessel not under command",
                    "A vessel of less than 20 meters in length",
                    "Another sailing vessel with the wind on the opposite side"
                ],
                'correct_index': 1,
                'marks': 2
            },
            {
                'question': f"Which of the following vessels must a vessel engaged in fishing keep out of the way of when underway?",
                'options': [
                    "A power-driven vessel",
                    "A vessel not under command",
                    "A sailing vessel",
                    "A vessel of less than 20 meters in length"
                ],
                'correct_index': 1,
                'marks': 2
            }
        ],
        '19': [
            {
                'question': f"In restricted visibility, what sound signal must a vessel underway and not making way give?",
                'options': [
                    "One short blast",
                    "Two prolonged blasts",
                    "One prolonged blast",
                    "Five short blasts"
                ],
                'correct_index': 1,
                'marks': 2
            },
            {
                'question': f"If, in restricted visibility you hear apparently forward of the beam a fog signal of another vessel, what action should you take?",
                'options': [
                    "Maintain course and speed",
                    "Reduce speed to minimum steerage and navigate with extreme caution",
                    "Come to a complete stop",
                    "Alter course to starboard regardless of the position of the other vessel"
                ],
                'correct_index': 1,
                'marks': 2
            }
        ],
        '21': [
            {
                'question': f"A power-driven vessel underway should exhibit which lights at night?",
                'options': [
                    "Green over red",
                    "Red over green",
                    "Green over white",
                    "Red over white"
                ],
                'correct_index': 2,
                'marks': 2
            }
        ],
        '23': [
            {
                'question': f"What lights must a power-driven vessel underway display?",
                'options': [
                    "Sidelights, masthead light, and sternlight",
                    "All-round white light only",
                    "Sidelights only",
                    "Flashing yellow light"
                ],
                'correct_index': 0,
                'marks': 2
            },
            {
                'question': f"When may a power-driven vessel display only sidelights and a sternlight?",
                'options': [
                    "When it is aground",
                    "When it is at anchor",
                    "When underway but not making way",
                    "When restricted in ability to manoeuvre"
                ],
                'correct_index': 2,
                'marks': 2
            }
        ],
        '28': [
            {
                'question': f"What does a vessel constrained by her draught exhibit?",
                'options': [
                    "Three red lights in a vertical line",
                    "One red light over two white lights",
                    "Three all-round red lights in a vertical line",
                    "Red over green"
                ],
                'correct_index': 2,
                'marks': 2
            }
        ],
        '30': [
            {
                'question': f"What lights does a vessel at anchor show at night?",
                'options': [
                    "One all-round white light where best seen",
                    "Two all-round white lights, one forward and one aft",
                    "One white light and one red light",
                    "No lights are required"
                ],
                'correct_index': 0,
                'marks': 2
            },
            {
                'question': f"What does a single black ball indicate during the day?",
                'options': [
                    "Vessel at anchor",
                    "Vessel not under command",
                    "Vessel towing",
                    "Vessel engaged in fishing"
                ],
                'correct_index': 0,
                'marks': 2
            }
        ],
        '32': [
            {
                'question': f"The term 'short blast' means a blast of:",
                'options': [
                    "About one second's duration",
                    "About two seconds' duration",
                    "About three seconds' duration",
                    "About five seconds' duration"
                ],
                'correct_index': 0,
                'marks': 2
            },
            {
                'question': f"The term 'prolonged blast' means a blast of:",
                'options': [
                    "About two seconds' duration",
                    "About three seconds' duration",
                    "From four to six seconds' duration",
                    "More than six seconds' duration"
                ],
                'correct_index': 2,
                'marks': 2
            }
        ],
        '34': [
            {
                'question': f"What sound signal shall a power-driven vessel make if altering course to starboard?",
                'options': [
                    "One short blast",
                    "Two short blasts",
                    "One prolonged blast",
                    "Three short blasts"
                ],
                'correct_index': 0,
                'marks': 2
            },
            {
                'question': f"What sound signal shall a power-driven vessel make if altering course to port?",
                'options': [
                    "One short blast",
                    "Two short blasts",
                    "One prolonged blast",
                    "Three short blasts"
                ],
                'correct_index': 1,
                'marks': 2
            },
            {
                'question': f"What is the meaning of three short blasts?",
                'options': [
                    "I am operating astern propulsion",
                    "I am altering course to starboard",
                    "I intend to overtake on starboard side",
                    "I am operating in fog"
                ],
                'correct_index': 0,
                'marks': 2
            },
            {
                'question': f"Which sound signal indicates a vessel is altering course to starboard?",
                'options': [
                    "One short blast",
                    "Two short blasts",
                    "Three short blasts",
                    "One prolonged blast"
                ],
                'correct_index': 0,
                'marks': 2
            }
        ],
        '35': [
            {
                'question': f"In restricted visibility, what sound signal should a power-driven vessel making way through the water give?",
                'options': [
                    "One short blast",
                    "Two prolonged blasts",
                    "One prolonged blast",
                    "Five short blasts"
                ],
                'correct_index': 2,
                'marks': 2
            },
            {
                'question': f"In an area of restricted visibility, which vessel would sound one prolonged and two short blasts at intervals not exceeding 2 minutes?",
                'options': [
                    "A vessel at anchor",
                    "A vessel not under command",
                    "A vessel restricted in her ability to maneuver",
                    "A vessel engaged in fishing"
                ],
                'correct_index': 2,
                'marks': 2
            },
            {
                'question': f"What sound signal should a vessel underway and making way in fog give?",
                'options': [
                    "One short blast every minute",
                    "One prolonged blast every two minutes",
                    "Two prolonged blasts every minute",
                    "Three short blasts every two minutes"
                ],
                'correct_index': 1,
                'marks': 2
            }
        ]
    }
    
    # Choose a rule-specific question if available, otherwise create a general question
    if rule['id'] in rule_questions:
        template = random.choice(rule_questions[rule['id']])
        
        # Validate data to prevent NoneType errors
        options = template.get('options', [])
        correct_index = template.get('correct_index', 0)
        marks = template.get('marks', 1)
        
        if not options or not isinstance(options, list) or correct_index >= len(options):
            raise ValueError(f"Invalid question structure for rule {rule['id']}")
            
        return {
            'type': 'Multiple Choice',
            'difficulty': difficulty,
            'question': template['question'],
            'choices': options,
            'answer': options[correct_index],
            'correct_index': correct_index,
            'marks': marks,
            'rule_reference': f"Rule {rule['id']} - {rule['title']}"
        }
    else:
        # Create a generic question if no specific template exists for this rule
        question = f"Which of the following statements about {rule['title']} is correct?"
        
        # Try to find relevant content for the correct answer
        content_parts = rule['content'].split('.')
        relevant_sentences = [s.strip() for s in content_parts if len(s.strip()) > 20 and len(s.strip()) < 150]
        
        if relevant_sentences:
            correct_option = random.choice(relevant_sentences)
        else:
            correct_option = f"This rule concerns {rule['title']}."
            
        # Create incorrect options
        incorrect_options = [
            f"This rule only applies to vessels over 50 meters in length.",
            f"This rule is only applicable in international waters.",
            f"This rule does not apply during daylight hours."
        ]
        
        options = [correct_option] + incorrect_options
        random.shuffle(options)
        correct_index = options.index(correct_option)
        
        return {
            'type': 'Multiple Choice',
            'difficulty': difficulty,
            'question': question,
            'choices': options,
            'answer': correct_option,
            'correct_index': correct_index,
            'marks': 2,
            'rule_reference': f"Rule {rule['id']} - {rule['title']}"
        }

def generate_exam_fill_blank(rule, difficulty):
    # Create exam-style fill in the blank questions matching the requested format
    rule_questions = {
        '3': [
            {
                'question': f"A vessel is said to be 'underway' when it is not at anchor, or made fast to the shore, or __________.",
                'answer': "aground",
                'marks': 1
            },
            {
                'question': f"Vessels shall be deemed to be in sight of one another only when one can be __________ from the other.",
                'answer': "observed visually",
                'marks': 1
            },
            {
                'question': f"The term 'vessel engaged in fishing' means any vessel fishing with nets, lines, trawls or other fishing apparatus which restrict __________.",
                'answer': "maneuverability",
                'marks': 1
            }
        ],
        '5': [
            {
                'question': f"Every vessel shall at all times maintain a proper look-out by sight and __________ as well as by all available means appropriate in the prevailing circumstances.",
                'answer': "hearing",
                'marks': 1
            }
        ],
        '6': [
            {
                'question': f"Every vessel shall at all times proceed at a __________ speed so that she can take proper and effective action to avoid collision.",
                'answer': "safe",
                'marks': 1
            }
        ],
        '7': [
            {
                'question': f"If there is any doubt about risk of collision, such risk shall be deemed to __________.",
                'answer': "exist",
                'marks': 1
            }
        ],
        '8': [
            {
                'question': f"Any action to avoid collision shall be taken in accordance with the Rules and shall be __________, made in ample time and with due regard to good seamanship.",
                'answer': "positive",
                'marks': 1
            }
        ],
        '9': [
            {
                'question': f"A vessel proceeding along the course of a narrow channel shall keep as near to the outer limit of the channel which lies on her __________ side as is safe and practicable.",
                'answer': "starboard",
                'marks': 1
            }
        ],
        '12': [
            {
                'question': f"When two sailing vessels are approaching with the wind on different sides, the vessel with the wind on the __________ side shall keep out of the way.",
                'answer': "port",
                'marks': 1
            }
        ],
        '13': [
            {
                'question': f"Any vessel overtaking any other shall keep out of the way of the vessel being __________.",
                'answer': "overtaken",
                'marks': 1
            },
            {
                'question': f"A vessel __________ the overtaking vessel must keep clear.",
                'answer': "being overtaken",
                'marks': 1
            }
        ],
        '14': [
            {
                'question': f"When two power-driven vessels are meeting head-on, each shall alter her course to __________.",
                'answer': "starboard",
                'marks': 1
            }
        ],
        '15': [
            {
                'question': f"When two power-driven vessels are crossing, the vessel which has the other on her __________ side shall keep out of the way.",
                'answer': "starboard",
                'marks': 1
            }
        ],
        '18': [
            {
                'question': f"A power-driven vessel underway shall keep out of the way of a __________ vessel.",
                'answer': "sailing",
                'marks': 1
            }
        ],
        '19': [
            {
                'question': f"In restricted visibility, every vessel shall proceed at a __________ speed adapted to the prevailing circumstances.",
                'answer': "safe",
                'marks': 1
            },
            {
                'question': f"In restricted visibility, a vessel must proceed at a safe __________.",
                'answer': "speed",
                'marks': 1
            }
        ],
        '21': [
            {
                'question': f"A vessel showing a green light on the starboard side and a red light on the __________ side is displaying proper sidelights.",
                'answer': "port",
                'marks': 1
            }
        ],
        '23': [
            {
                'question': f"A power-driven vessel underway shall exhibit sidelights and a __________.",
                'answer': "sternlight",
                'marks': 1
            }
        ],
        '27': [
            {
                'question': f"A vessel not under command shall exhibit two all-round __________ lights in a vertical line where they can best be seen.",
                'answer': "red",
                'marks': 1
            },
            {
                'question': f"The day shape for a vessel restricted in her ability to maneuver is a ball, __________, ball.",
                'answer': "diamond",
                'marks': 1
            }
        ],
        '28': [
            {
                'question': f"The day shape for a vessel constrained by her draft is a __________.",
                'answer': "cylinder",
                'marks': 1
            }
        ],
        '30': [
            {
                'question': f"A vessel at anchor shall exhibit where it can best be seen an all-round __________ light.",
                'answer': "white",
                'marks': 1
            },
            {
                'question': f"A vessel at anchor shall show an all-round white light where it can best be __________.",
                'answer': "seen",
                'marks': 1
            }
        ],
        '34': [
            {
                'question': f"One short blast means 'I am altering my course to __________'.",
                'answer': "starboard",
                'marks': 1
            },
            {
                'question': f"The sound signal for altering course to port is __________.",
                'answer': "two short blasts",
                'marks': 1
            },
            {
                'question': f"The sound signal for altering course to starboard is __________.",
                'answer': "one short blast",
                'marks': 1
            },
            {
                'question': f"The sound signal for operating astern propulsion is __________.",
                'answer': "three short blasts",
                'marks': 1
            }
        ],
        '35': [
            {
                'question': f"A power-driven vessel making way through the water in restricted visibility shall sound at intervals of not more than 2 minutes one __________ blast.",
                'answer': "prolonged",
                'marks': 1
            },
            {
                'question': f"A power-driven vessel underway but stopped and making no way through the water shall sound at intervals of not more than 2 minutes __________.",
                'answer': "two prolonged blasts",
                'marks': 1
            }
        ]
    }
    
    # Choose a rule-specific question if available, otherwise create a general question
    if rule['id'] in rule_questions:
        template = random.choice(rule_questions[rule['id']])
        
        if not template.get('question') or not template.get('answer'):
            raise ValueError(f"Invalid template for rule {rule['id']}")
            
        return {
            'type': 'Fill in the Gap',
            'difficulty': difficulty,
            'question': template['question'],
            'answer': template['answer'],
            'marks': template.get('marks', 1),
            'rule_reference': f"Rule {rule['id']} - {rule['title']}"
        }
    else:
        # Create a generic fill-in-the-blank question if no specific template exists
        return create_fallback_fill_blank(rule, difficulty)

def generate_exam_written_answer(rule, difficulty):
    # Create exam-style written answer questions with direct, specific questions matching the requested format
    rule_questions = {
        '1': [
            {
                'question': f"Where do the Collision Regulations apply? 2 Marks",
                'answer': "To all vessels upon the high seas and in all waters connected therewith navigable by seagoing vessels.",
                'marks': 2
            }
        ],
        '3': [
            {
                'question': f"Define a 'vessel not under command' as per the COLREGS. 2 Marks",
                'answer': "A vessel which through some exceptional circumstance is unable to maneuver as required by these Rules and is therefore unable to keep out of the way of another vessel.",
                'marks': 2
            },
            {
                'question': f"Define a 'vessel restricted in her ability to maneuver' as per the COLREGS. 2 Marks",
                'answer': "A vessel which from the nature of her work is restricted in her ability to maneuver as required by these Rules and is therefore unable to keep out of the way of another vessel.",
                'marks': 2
            },
            {
                'question': f"List three types of vessels that would be classified as 'restricted in their ability to maneuver'. 3 Marks",
                'answer': "1. A vessel engaged in laying, servicing or picking up a navigation mark, submarine cable or pipeline\n2. A vessel engaged in dredging, surveying or underwater operations\n3. A vessel engaged in replenishment or transferring persons, provisions or cargo while underway",
                'marks': 3
            }
        ],
        '5': [
            {
                'question': f"What is required to maintain a proper look-out according to Rule 5? 2 Marks",
                'answer': "Every vessel shall at all times maintain a proper look-out by sight and hearing as well as by all available means appropriate in the prevailing circumstances and conditions so as to make a full appraisal of the situation and of the risk of collision.",
                'marks': 2
            }
        ],
        '6': [
            {
                'question': f"List four factors that should be taken into account when determining a safe speed. 4 Marks",
                'answer': "1. The state of visibility\n2. The traffic density including concentrations of fishing vessels or any other vessels\n3. The maneuverability of the vessel with special reference to stopping distance and turning ability\n4. The state of wind, sea and current, and the proximity of navigational hazards",
                'marks': 4
            }
        ],
        '7': [
            {
                'question': f"When should you assume that risk of collision exists? 1 Mark",
                'answer': "If there is any doubt, then risk of collision shall be deemed to exist.",
                'marks': 1
            }
        ],
        '8': [
            {
                'question': f"What three characteristics should any action taken to avoid collision have? 3 Marks",
                'answer': "1. It shall be positive\n2. It shall be made in ample time\n3. It shall be made with due regard to the observance of good seamanship",
                'marks': 3
            }
        ],
        '9': [
            {
                'question': f"When navigating in a narrow channel, on which side should a vessel keep? 2 Marks",
                'answer': "A vessel proceeding along the course of a narrow channel or fairway shall keep as near to the outer limit of the channel or fairway which lies on her starboard side as is safe and practicable.",
                'marks': 2
            }
        ],
        '10': [
            {
                'question': f"What action should a vessel take according to Rule 10 if it is NOT using a traffic separation scheme? 2 Marks",
                'answer': "A vessel not using a traffic separation scheme shall avoid it by as wide a margin as is practicable.",
                'marks': 2
            },
            {
                'question': f"How should vessels proceed when in a traffic lane? 2 Marks",
                'answer': "Vessels shall proceed in the appropriate traffic lane in the general direction of traffic flow for that lane.",
                'marks': 2
            },
            {
                'question': f"What is the proper way to cross a traffic lane? 2 Marks",
                'answer': "A vessel crossing a traffic lane shall do so on a heading as nearly as practicable at right angles to the general direction of traffic flow.",
                'marks': 2
            },
            {
                'question': f"Which vessels should avoid the separation zones in a traffic separation scheme? 2 Marks",
                'answer': "A vessel shall so far as practicable avoid crossing traffic lanes, but if obliged to do so shall cross on a heading as nearly as practicable at right angles to the general direction of traffic flow.",
                'marks': 2
            }
        ],
        '11': [
            {
                'question': f"When do the Rules of Section I of Part B apply? 2 Marks",
                'answer': "The Rules in Section I apply to vessels in sight of one another.",
                'marks': 2
            },
            {
                'question': f"What is meant by 'vessels in sight of one another'? 2 Marks",
                'answer': "Vessels are in sight of one another only when one can be observed visually from the other.",
                'marks': 2
            }
        ],
        '12': [
            {
                'question': f"Name two situations where one sailing vessel must keep out of the way of another sailing vessel. 2 Marks",
                'answer': "1. When each has the wind on a different side, the vessel which has the wind on the port side shall keep out of the way\n2. When both have the wind on the same side, the vessel which is to windward shall keep out of the way of the vessel which is to leeward",
                'marks': 2
            }
        ],
        '13': [
            {
                'question': f"How is an overtaking vessel defined in Rule 13? 2 Marks",
                'answer': "A vessel shall be deemed to be overtaking when coming up with another vessel from a direction more than 22.5 degrees abaft her beam.",
                'marks': 2
            }
        ],
        '14': [
            {
                'question': f"What action should be taken when two power-driven vessels are meeting head-on? 2 Marks",
                'answer': "Each vessel shall alter her course to starboard so that each shall pass on the port side of the other.",
                'marks': 2
            }
        ],
        '15': [
            {
                'question': f"In a crossing situation between two power-driven vessels, which vessel should keep out of the way? 2 Marks",
                'answer': "The vessel which has the other on her own starboard side shall keep out of the way and shall, if the circumstances of the case admit, avoid crossing ahead of the other vessel.",
                'marks': 2
            }
        ],
        '16': [
            {
                'question': f"What action must a give-way vessel take? 2 Marks",
                'answer': "Every vessel which is directed to keep out of the way of another vessel shall, so far as possible, take early and substantial action to keep well clear.",
                'marks': 2
            },
            {
                'question': f"How should a give-way vessel maneuver to keep out of the way of another vessel? 2 Marks",
                'answer': "It should take early and substantial action to keep well clear.",
                'marks': 2
            }
        ],
        '17': [
            {
                'question': f"What is the meaning of the term 'stand-on vessel'? 2 Marks",
                'answer': "A vessel that must maintain course and speed when one of two vessels is to keep out of the way of the other.",
                'marks': 2
            }
        ],
        '18': [
            {
                'question': f"Name two situations where a sailing vessel must keep out of the way of another vessel. 2 Marks",
                'answer': "1. A vessel not under command\n2. A vessel restricted in her ability to maneuver",
                'marks': 2
            }
        ],
        '19': [
            {
                'question': f"What action should a vessel take if it hears a fog signal apparently forward of the beam in restricted visibility? 2 Marks",
                'answer': "Reduce speed to the minimum at which she can be kept on her course. If necessary take all her way off and in any event navigate with extreme caution until danger of collision is over.",
                'marks': 2
            },
            {
                'question': f"What action should a vessel take when it hears a fog signal forward of the beam? 2 Marks",
                'answer': "Reduce speed and navigate with extreme caution until the danger of collision is over.",
                'marks': 2
            }
        ],
        '22': [
            {
                'question': f"What is the minimum visibility range for sidelights on vessels of 50 meters or more in length? 2 Marks",
                'answer': "3 miles (or 6 kilometers).",
                'marks': 2
            },
            {
                'question': f"State the visibility range requirements for masthead lights on vessels less than 50 meters in length. 2 Marks",
                'answer': "At least 3 miles (6 kilometers), except that where the length of the vessel is less than 12 meters, at least 2 miles (4 kilometers).",
                'marks': 2
            }
        ],
        '23': [
            {
                'question': f"What lights must a power-driven vessel underway display? 3 Marks",
                'answer': "1. A masthead light forward\n2. A second masthead light abaft of and higher than the forward one (if 50m or more in length)\n3. Sidelights\n4. A sternlight",
                'marks': 3
            }
        ],
        '24': [
            {
                'question': f"What lights must a power-driven vessel towing another vessel show when the length of the tow exceeds 200 meters? 2 Marks",
                'answer': "Three masthead lights in a vertical line, sidelights, a sternlight, and a towing light above the sternlight.",
                'marks': 2
            },
            {
                'question': f"What day shape must a vessel being towed display if the length of the tow exceeds 200 meters? 2 Marks",
                'answer': "A diamond shape where it can best be seen.",
                'marks': 2
            }
        ],
        '29': [
            {
                'question': f"What lights must a pilot vessel on duty display at night? 2 Marks",
                'answer': "At or near the masthead, two all-round lights in a vertical line, the upper being white and the lower red, and when underway, also sidelights and a sternlight.",
                'marks': 2
            },
            {
                'question': f"What day shape does a pilot vessel on duty display? 1 Mark",
                'answer': "A pilot flag 'H' (Hotel - white and red flag).",
                'marks': 1
            }
        ],
        '30': [
            {
                'question': f"What lights does a vessel at anchor show at night? 2 Marks",
                'answer': "One all-round white light where best seen; two if over 50m (one forward, one aft at a lower level).",
                'marks': 2
            },
            {
                'question': f"Name the day shape shown by a vessel aground. 1 Mark",
                'answer': "Three black balls in a vertical line.",
                'marks': 1
            }
        ],
        '31': [
            {
                'question': f"What should a seaplane do when it cannot comply fully with the rules for lights? 2 Marks",
                'answer': "It shall exhibit lights and shapes as closely similar in characteristics and position as is possible.",
                'marks': 2
            },
            {
                'question': f"What are seaplanes required to do when maneuvering on water? 2 Marks",
                'answer': "Comply with the rules for vessels as far as is practicable.",
                'marks': 2
            }
        ],
        '32': [
            {
                'question': f"How long is a 'short blast'? 1 Mark",
                'answer': "About one second's duration.",
                'marks': 1
            },
            {
                'question': f"How long is a 'prolonged blast'? 1 Mark",
                'answer': "From four to six seconds' duration.",
                'marks': 1
            },
            {
                'question': f"Define the term 'prolonged blast' as specified in Rule 32. 2 Marks",
                'answer': "A blast of from four to six seconds' duration.",
                'marks': 2
            }
        ],
        '33': [
            {
                'question': f"What equipment for sound signals must vessels of 12 meters or more in length carry? 2 Marks",
                'answer': "A whistle and a bell. Vessels of 100 meters or more in length shall, in addition, be provided with a gong.",
                'marks': 2
            },
            {
                'question': f"What sound signal equipment must vessels less than 12 meters in length carry? 2 Marks",
                'answer': "Some means of making an efficient sound signal, though they are not obliged to carry a whistle, bell or gong.",
                'marks': 2
            }
        ],
        '34': [
            {
                'question': f"What sound signal should a power-driven vessel make when altering course to starboard? 1 Mark",
                'answer': "One short blast.",
                'marks': 1
            },
            {
                'question': f"What sound signal should a power-driven vessel make when operating astern propulsion? 1 Mark",
                'answer': "Three short blasts.",
                'marks': 1
            },
            {
                'question': f"What is the danger signal? 1 Mark",
                'answer': "Five or more short and rapid blasts on the whistle.",
                'marks': 1
            },
            {
                'question': f"What action should a vessel take when in doubt of another vessel's intentions? 1 Mark",
                'answer': "Give at least five short and rapid blasts on the whistle.",
                'marks': 1
            }
        ],
        '35': [
            {
                'question': f"What sound signal should a power-driven vessel give when operating in restricted visibility and making way through the water? 2 Marks",
                'answer': "One prolonged blast at intervals of not more than 2 minutes.",
                'marks': 2
            },
            {
                'question': f"What sound signal should a vessel at anchor make in restricted visibility? 2 Marks",
                'answer': "Ring the bell rapidly for about 5 seconds at intervals of not more than one minute.",
                'marks': 2
            },
            {
                'question': f"What sound signal should a vessel underway and making way in fog give? 2 Marks",
                'answer': "One prolonged blast every two minutes.",
                'marks': 2
            }
        ]
    }
    
    # Choose a rule-specific question if available, otherwise create a specific question
    if rule['id'] in rule_questions:
        template = random.choice(rule_questions[rule['id']])
        
        if not template.get('question') or not template.get('answer') or not template.get('marks'):
            raise ValueError(f"Invalid question structure for rule {rule['id']}")
            
        return {
            'type': 'Written Answer',
            'difficulty': difficulty,
            'question': template['question'],
            'answer': template['answer'],
            'rule_reference': f"Rule {rule['id']} - {rule['title']}",
            'marks': template['marks']
        }
    else:
        # Create a specific question using content from the rule
        return create_fallback_written_answer(rule, difficulty)

def create_fallback_fill_blank(rule, difficulty):
    """Create a simple fallback fill-in-the-blank question when normal generation fails"""
    return {
        'type': 'Fill in the Gap',
        'difficulty': difficulty,
        'question': f"According to Rule {rule['id']}, ships must follow __________ procedures when navigating.",
        'answer': "specific",
        'rule_reference': f"Rule {rule['id']} - {rule['title']}",
        'marks': 1
    }

def create_fallback_written_answer(rule, difficulty):
    """Create a specific fallback written answer question when no template exists"""
    
    # Check rule content for specific keywords to create targeted questions
    content = rule['content'].lower()
    rule_title = rule['title'].lower()
    
    # Rule-specific fallback formats based on rule title and content
    if "equipment" in rule_title and "sound" in rule_title:
        question = f"What sound signal equipment is required on vessels according to Rule {rule['id']}? 2 Marks"
        answer = get_relevant_content_snippet(rule['content'], 150)
    elif "light" in content or "display" in content or "exhibit" in content:
        question = f"List the light requirements specified in Rule {rule['id']}. 2 Marks"
        answer = "The light requirements include: " + get_relevant_content_snippet(rule['content'], 150)
    elif ("sound" in content or "signal" in content) and not "equipment" in rule_title:
        question = f"What sound signals are required according to Rule {rule['id']}? 2 Marks"
        answer = "The required sound signals are: " + get_relevant_content_snippet(rule['content'], 150)
    elif "day" in content and "shape" in content:
        question = f"What day shapes are required according to Rule {rule['id']}? 2 Marks"
        answer = "The day shapes required are: " + get_relevant_content_snippet(rule['content'], 150)
    elif "action" in content or "maneuver" in content or "manoeuvre" in content:
        question = f"What action should a vessel take according to Rule {rule['id']}? 2 Marks"
        answer = "The vessel should: " + get_relevant_content_snippet(rule['content'], 150)
    else:
        # Extract a specific requirement to quiz on
        content_parts = rule['content'].split('.')
        relevant_parts = [p.strip() for p in content_parts if len(p.strip()) > 20 and len(p.strip()) < 200]
        
        if relevant_parts:
            specific_content = random.choice(relevant_parts)
            # Create a question about a specific aspect rather than the entire rule
            question = f"According to Rule {rule['id']}, explain {specific_content[:30]}... 2 Marks"
            answer = specific_content
        else:
            # Last resort specific question
            question = f"List two key points from Rule {rule['id']} regarding {rule['title']}. 2 Marks"
            answer = "Key points from Rule {rule['id']}: " + get_relevant_content_snippet(rule['content'], 150)
    
    return {
        'type': 'Written Answer',
        'difficulty': difficulty,
        'question': question,
        'answer': answer,
        'rule_reference': f"Rule {rule['id']} - {rule['title']}",
        'marks': 2
    }

def get_relevant_content_snippet(content, max_length=150):
    """Extract a relevant piece of content from the rule text"""
    content_parts = content.split('.')
    relevant_parts = [p.strip() for p in content_parts if len(p.strip()) > 20 and len(p.strip()) < max_length]
    
    if relevant_parts:
        return random.choice(relevant_parts)
    else:
        # If we can't find a good snippet, return a truncated version of the content
        return content[:max_length] + "..."

def generate_true_false(rule, difficulty):
    """Generate true/false questions about the IRPCS rules"""
    # Dictionary of rule-specific true/false questions
    rule_questions = {
        '1': [
            {
                'question': "The Collision Regulations apply only to vessels in international waters.",
                'answer': False,
                'marks': 1,
                'explanation': "The Collision Regulations apply to all vessels upon the high seas and in all waters connected therewith navigable by seagoing vessels."
            }
        ],
        '3': [
            {
                'question': "A vessel under sail and power must be treated as a sailing vessel.",
                'answer': False,
                'marks': 1,
                'explanation': "A vessel under sail and power is considered a power-driven vessel, not a sailing vessel."
            },
            {
                'question': "The term 'vessel engaged in fishing' includes any vessel fishing with trolling lines.",
                'answer': False,
                'marks': 1,
                'explanation': "The term does not include vessels fishing with trolling lines as they do not restrict maneuverability."
            },
            {
                'question': "A sailing vessel under power is treated as a power-driven vessel.",
                'answer': True,
                'marks': 1,
                'explanation': "When a sailing vessel is using its engine, it is treated as a power-driven vessel, not a sailing vessel."
            }
        ],
        '5': [
            {
                'question': "Maintaining a proper look-out involves both sight and hearing.",
                'answer': True,
                'marks': 1,
                'explanation': "Rule 5 states that every vessel shall at all times maintain a proper look-out by sight and hearing."
            }
        ],
        '6': [
            {
                'question': "The state of visibility is a factor in determining safe speed.",
                'answer': True,
                'marks': 1,
                'explanation': "The state of visibility is one of the factors listed in Rule 6 that should be taken into account when determining a safe speed."
            }
        ],
        '7': [
            {
                'question': "Vessels are in sight of one another when they can be detected on radar.",
                'answer': False,
                'marks': 1,
                'explanation': "Vessels are in sight of one another only when one can be observed visually from the other."
            }
        ],
        '9': [
            {
                'question': "In a crossing situation, the vessel on the port side must give way.",
                'answer': True,
                'marks': 1,
                'explanation': "In a crossing situation, the vessel which has the other on her starboard side shall keep out of the way."
            }
        ],
        '12': [
            {
                'question': "When two sailing vessels have the wind on the same side, the vessel which is to leeward shall keep out of the way.",
                'answer': False,
                'marks': 1,
                'explanation': "The vessel which is to windward shall keep out of the way of the vessel which is to leeward."
            }
        ],
        '13': [
            {
                'question': "An overtaking vessel remains an overtaking vessel even if the bearing changes.",
                'answer': True,
                'marks': 1,
                'explanation': "Any subsequent alteration of the bearing between the two vessels shall not make the overtaking vessel a crossing vessel."
            }
        ],
        '14': [
            {
                'question': "In a head-on situation, both vessels must alter course to port.",
                'answer': False,
                'marks': 1,
                'explanation': "Both vessels must alter course to starboard."
            }
        ],
        '15': [
            {
                'question': "In a crossing situation, the vessel that has the other on her port side must keep out of the way.",
                'answer': False,
                'marks': 1,
                'explanation': "The vessel that has the other on her starboard side must keep out of the way."
            }
        ],
        '18': [
            {
                'question': "A power-driven vessel must keep out of the way of a sailing vessel.",
                'answer': True,
                'marks': 1,
                'explanation': "Rule 18 states that a power-driven vessel underway shall keep out of the way of a sailing vessel."
            },
            {
                'question': "A power-driven vessel must always give way to a sailing vessel.",
                'answer': False,
                'marks': 1,
                'explanation': "There are exceptions, such as when the sailing vessel is overtaking or when the power-driven vessel is not under command."
            }
        ],
        '19': [
            {
                'question': "In restricted visibility, a vessel should always maintain its normal speed.",
                'answer': False,
                'marks': 1,
                'explanation': "A vessel should proceed at a safe speed adapted to the prevailing circumstances."
            },
            {
                'question': "In restricted visibility, vessels must sound appropriate signals.",
                'answer': True,
                'marks': 1,
                'explanation': "Rule 35 requires vessels to sound specified signals in or near areas of restricted visibility."
            }
        ],
        '23': [
            {
                'question': "A power-driven vessel of less than 50 meters in length must show two masthead lights.",
                'answer': False,
                'marks': 1,
                'explanation': "A vessel of less than 50 meters in length is not obliged to exhibit a second masthead light."
            }
        ],
        '28': [
            {
                'question': "A vessel constrained by her draft shows three black balls by day.",
                'answer': False,
                'marks': 1,
                'explanation': "A vessel constrained by her draft may exhibit a cylinder by day."
            }
        ],
        '30': [
            {
                'question': "A vessel at anchor must exhibit a sternlight at night.",
                'answer': False,
                'marks': 1,
                'explanation': "A vessel at anchor shows an all-round white light where it can best be seen; vessels over 50m show a second all-round white light at the stern at a lower level."
            },
            {
                'question': "A vessel at anchor should exhibit an all-round white light.",
                'answer': True,
                'marks': 1,
                'explanation': "A vessel at anchor shall exhibit an all-round white light where it can best be seen."
            }
        ],
        '34': [
            {
                'question': "Sound signals are only required during the day.",
                'answer': False,
                'marks': 1,
                'explanation': "Sound signals are required whenever necessary for safe navigation, day or night."
            }
        ],
        '35': [
            {
                'question': "A vessel at anchor in restricted visibility must sound a bell rapidly for about 5 seconds at intervals of not more than one minute.",
                'answer': True,
                'marks': 1,
                'explanation': "Rule 35(g) requires a vessel at anchor to ring the bell rapidly for about 5 seconds at intervals of not more than one minute."
            }
        ]
    }
    
    # General questions that apply to multiple rules
    general_questions = [
        {
            'question': "A vessel under sail and power must be treated as a sailing vessel.",
            'answer': False,
            'marks': 1,
            'explanation': "A vessel under sail and power is considered a power-driven vessel, not a sailing vessel."
        },
        {
            'question': "A sailing vessel under power is treated as a power-driven vessel.",
            'answer': True,
            'marks': 1,
            'explanation': "When a sailing vessel is using its engine, it is treated as a power-driven vessel, not a sailing vessel."
        },
        {
            'question': "A power-driven vessel must always give way to a sailing vessel.",
            'answer': False,
            'marks': 1,
            'explanation': "There are exceptions, such as when the sailing vessel is overtaking or when the power-driven vessel is not under command."
        },
        {
            'question': "In restricted visibility, vessels must sound appropriate signals.",
            'answer': True,
            'marks': 1,
            'explanation': "Rule 35 requires vessels to sound specified signals in or near areas of restricted visibility."
        },
        {
            'question': "A vessel at anchor should exhibit an all-round white light.",
            'answer': True,
            'marks': 1,
            'explanation': "A vessel at anchor shall exhibit an all-round white light where it can best be seen."
        }
    ]
    
    # Choose a rule-specific question if available, otherwise use a general question
    if rule['id'] in rule_questions:
        template = random.choice(rule_questions[rule['id']])
    else:
        # Use a general question if no specific template exists for this rule
        template = random.choice(general_questions)
        
    if 'question' not in template or 'answer' not in template:
        raise ValueError(f"Invalid true/false question structure")
        
    result = {
        'type': 'True/False',
        'difficulty': difficulty,
        'question': template['question'],
        'answer': "True" if template['answer'] else "False",
        'marks': template.get('marks', 1),
        'rule_reference': f"Rule {rule['id']} - {rule['title']}"
    }
    
    if 'explanation' in template:
        result['explanation'] = template['explanation']
        
    return result

if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5000, debug=True) 